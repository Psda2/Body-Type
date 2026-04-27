import os
from docx import Document
from docx.shared import Pt
import re

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.readlines()

    in_table = False
    table_data = []

    for line in content:
        line = line.strip('\n')
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            continue

        # Tables
        if '|' in line and '-' * 3 in line and not any(c.isalnum() for c in line):
             # This is the separator line, skip it but we know we are in a table
             in_table = True
             continue
        
        if '|' in line:
            cells = [c.strip() for c in line.split('|') if c.strip() or (line.startswith('|') and line.endswith('|'))]
            # Clean up empty strings at ends if they exist due to leading/trailing pipes
            if line.strip().startswith('|'): cells = [c.strip() for c in line.strip('|').split('|')]
            
            if not in_table:
                # Potential header row
                table_data = [cells]
                in_table = True
            else:
                table_data.append(cells)
            continue
        elif in_table:
            # End of table
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                if cols > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            if j < cols:
                                # Strip Markdown bold/italic from table cells
                                clean_cell = re.sub(r'\*\*|\*', '', cell_data)
                                table.cell(i, j).text = clean_cell
            table_data = []
            in_table = False

        # Bold and Lists
        if line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            # Rudimentary bold handling: **text**
            parts = re.split(r'(\*\*.*?\*\*)', line[2:])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            continue

        # Normal paragraphs
        if line.strip():
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Strip other MD like [link](url)
                    clean_part = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', part)
                    # Strip LaTeX $...$
                    clean_part = re.sub(r'\$(.*?)\$', r'\1', clean_part)
                    p.add_run(clean_part)
        else:
            # Empty line
            pass

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    MD_FILE = os.path.join(BASE_DIR, "../thesis_project_report.md")
    DOCX_FILE = os.path.join(BASE_DIR, "../thesis_project_report.docx")
    convert_md_to_docx(MD_FILE, DOCX_FILE)
